#!/usr/bin/env python3
"""
Comprehensive Unit Tests for CVCustomizer Class
Each test is designed to be self-running and independent.
Supports parameterized testing with default values.
"""

import unittest
import tempfile
import os
import yaml
import json
import argparse
import shutil
from unittest.mock import Mock, patch, mock_open, MagicMock
from docx import Document
import requests
from cv_customizer import CVCustomizer

# Default test parameters
DEFAULT_TEST_PARAMS = {
    'api_key': 'test-api-key-12345',
    'api_endpoint': 'https://test-api.openai.com/v1/chat/completions',
    'config_file': 'prompts_config.yaml',
    'test_yaml_file': 'test_cv_template.yaml',
    'test_output_file': 'test_output.docx',
    'test_job_description': 'Looking for a software engineer with Python experience',
    'test_company_name': 'Tech Corp',
    'test_job_title': 'Senior Software Engineer',
    'test_name': 'John Doe',
    'test_email': 'john.doe@example.com',
    'test_phone': '123-456-7890',
    'test_linkedin': 'linkedin.com/in/johndoe',
    'test_nationality': 'American'
}


class ParameterizedTestCase(unittest.TestCase):
    """Base class for parameterized tests"""
    
    @classmethod
    def setUpClass(cls):
        """Set up test parameters"""
        cls.test_params = DEFAULT_TEST_PARAMS.copy()
        
        # Allow overriding parameters via command line or environment
        parser = argparse.ArgumentParser(description='CVCustomizer Unit Tests')
        parser.add_argument('--api-key', default=cls.test_params['api_key'])
        parser.add_argument('--api-endpoint', default=cls.test_params['api_endpoint'])
        parser.add_argument('--test-yaml-file', default=cls.test_params['test_yaml_file'])
        parser.add_argument('--test-output-file', default=cls.test_params['test_output_file'])
        parser.add_argument('--test-job-description', default=cls.test_params['test_job_description'])
        parser.add_argument('--test-company-name', default=cls.test_params['test_company_name'])
        parser.add_argument('--test-job-title', default=cls.test_params['test_job_title'])
        parser.add_argument('--test-name', default=cls.test_params['test_name'])
        parser.add_argument('--test-email', default=cls.test_params['test_email'])
        parser.add_argument('--test-phone', default=cls.test_params['test_phone'])
        parser.add_argument('--test-linkedin', default=cls.test_params['test_linkedin'])
        parser.add_argument('--test-nationality', default=cls.test_params['test_nationality'])
        parser.add_argument('--config-file', default=cls.test_params['config_file'])
        
        # Parse arguments (ignore unknown args for individual test runs)
        args, unknown = parser.parse_known_args()
        
        # Update test parameters with provided values
        cls.test_params.update({
            'api_key': args.api_key,
            'api_endpoint': args.api_endpoint,
            'config_file': args.config_file,
            'test_yaml_file': args.test_yaml_file,
            'test_output_file': args.test_output_file,
            'test_job_description': args.test_job_description,
            'test_company_name': args.test_company_name,
            'test_job_title': args.test_job_title,
            'test_name': args.test_name,
            'test_email': args.test_email,
            'test_phone': args.test_phone,
            'test_linkedin': args.test_linkedin,
            'test_nationality': args.test_nationality
        })
    
    def get_test_param(self, param_name, default=None):
        """Get a test parameter with fallback to default"""
        return self.test_params.get(param_name, default)
    
    def save_temp_file_copy(self, temp_file_path, file_type="output", verbose=True):
        """Save a copy of temporary file in local directory for debugging"""
        local_copy_path = f"test_{file_type}_{os.getpid()}.{temp_file_path.split('.')[-1]}"
        shutil.copy2(temp_file_path, local_copy_path)
        if verbose:
            print(f"Test {file_type} saved to: {local_copy_path}")
        return local_copy_path


class TestCVCustomizerInit(ParameterizedTestCase):
    """Test cases for CVCustomizer.__init__ method"""
    
    def test_init_with_api_key_and_endpoint(self):
        """Test initialization with provided API key and endpoint"""
        api_key = self.get_test_param('api_key', "test-api-key")
        api_endpoint = self.get_test_param('api_endpoint', "https://test-api.com/v1/chat/completions")
        config_file = self.get_test_param('config_file', "prompts_config.yaml")
        
        customizer = CVCustomizer(api_key, api_endpoint, config_file)
        
        self.assertEqual(customizer.api_key, api_key)
        self.assertEqual(customizer.api_endpoint, api_endpoint)
        self.assertIsNotNone(customizer.config)
    
    def test_init_with_none_values(self):
        """Test initialization with None values"""
        with patch.dict(os.environ, {'LLM_API_KEY': 'env-key', 'LLM_API_ENDPOINT': 'env-endpoint'}):
            config_file = self.get_test_param('config_file', "prompts_config.yaml")
            customizer = CVCustomizer(None, None, config_file)
            
            self.assertEqual(customizer.api_key, 'env-key')
            self.assertEqual(customizer.api_endpoint, 'env-endpoint')
            self.assertIsNotNone(customizer.config)
    
    def test_init_with_no_api_key(self):
        """Test initialization with no API key"""
        with patch.dict(os.environ, {}, clear=True):
            with patch('builtins.print') as mock_print:
                config_file = self.get_test_param('config_file', "prompts_config.yaml")
                customizer = CVCustomizer(None, None, config_file)
                
                self.assertIsNone(customizer.api_key)
                self.assertEqual(customizer.api_endpoint, "https://api.openai.com/v1/chatP/completions")
                self.assertIsNotNone(customizer.config)
                mock_print.assert_called_once()


class TestCVCustomizerLoadCVTemplate(ParameterizedTestCase):
    """Test cases for CVCustomizer.load_cv_template method"""
    
    def setUp(self):
        """Set up test fixtures"""
        api_key = self.get_test_param('api_key', "test-key")
        api_endpoint = self.get_test_param('api_endpoint', "test-endpoint")
        config_file = self.get_test_param('config_file', "prompts_config.yaml")
        self.customizer = CVCustomizer(api_key, api_endpoint, config_file)
        
        self.sample_cv_data = {
            "name": self.get_test_param('test_name', "John Doe"),
            "contact": {"email": self.get_test_param('test_email', "john@example.com")},
            "experience": []
        }
    
    def test_load_cv_template_success(self):
        """Test successful loading of CV template"""
        with tempfile.NamedTemporaryFile(mode='w', suffix='.yaml', delete=False) as temp_file:
            yaml.dump(self.sample_cv_data, temp_file)
            temp_file.flush()
            
            try:
                result = self.customizer.load_cv_template(temp_file.name)
                self.assertEqual(result, self.sample_cv_data)
                
                # Save a copy in local directory for debugging (without printing)
                self.save_temp_file_copy(temp_file.name, "template", verbose=False)
                
            finally:
                os.unlink(temp_file.name)
    
    def test_load_cv_template_file_not_found(self):
        """Test loading non-existent file"""
        with self.assertRaises(FileNotFoundError):
            self.customizer.load_cv_template("non_existent_file.yaml")
    
    def test_load_cv_template_invalid_yaml(self):
        """Test loading invalid YAML file"""
        with tempfile.NamedTemporaryFile(mode='w', suffix='.yaml', delete=False) as temp_file:
            temp_file.write("invalid: yaml: content: [")
            temp_file.flush()
            
            try:
                with self.assertRaises(ValueError):
                    self.customizer.load_cv_template(temp_file.name)
            finally:
                os.unlink(temp_file.name)


class TestCVCustomizerCallLLMAPI(ParameterizedTestCase):
    """Test cases for CVCustomizer.call_llm_api method"""
    
    def setUp(self):
        """Set up test fixtures"""
        api_key = self.get_test_param('api_key', "test-key")
        api_endpoint = self.get_test_param('api_endpoint', "test-endpoint")
        config_file = self.get_test_param('config_file', "prompts_config.yaml")
        self.customizer = CVCustomizer(api_key, api_endpoint, config_file)
    
    def test_call_llm_api_no_api_key(self):
        """Test API call with no API key"""
        customizer = CVCustomizer(None, None)
        
        with patch('builtins.print') as mock_print:
            result = customizer.call_llm_api("test prompt")
            
            self.assertEqual(result, "test prompt")
            mock_print.assert_called_once()
    
    @patch('requests.post')
    def test_call_llm_api_success(self, mock_post):
        """Test successful API call"""
        mock_response = Mock()
        mock_response.json.return_value = {
            "choices": [{"message": {"content": "Generated response"}}]
        }
        mock_response.raise_for_status.return_value = None
        mock_post.return_value = mock_response
        
        result = self.customizer.call_llm_api("test prompt", "system message")
        
        self.assertEqual(result, "Generated response")
        mock_post.assert_called_once()
    
    @patch('requests.post')
    def test_call_llm_api_without_system_message(self, mock_post):
        """Test API call without system message"""
        mock_response = Mock()
        mock_response.json.return_value = {
            "choices": [{"message": {"content": "Generated response"}}]
        }
        mock_response.raise_for_status.return_value = None
        mock_post.return_value = mock_response
        
        result = self.customizer.call_llm_api("test prompt")
        
        self.assertEqual(result, "Generated response")
        call_args = mock_post.call_args
        self.assertEqual(len(call_args[1]['json']['messages']), 1)
        self.assertEqual(call_args[1]['json']['messages'][0]['role'], 'user')
    
    @patch('requests.post')
    def test_call_llm_api_unexpected_response(self, mock_post):
        """Test API call with unexpected response format"""
        mock_response = Mock()
        mock_response.json.return_value = {"unexpected": "format"}
        mock_response.raise_for_status.return_value = None
        mock_post.return_value = mock_response
        
        with patch('builtins.print') as mock_print:
            result = self.customizer.call_llm_api("test prompt")
            
            self.assertEqual(result, "test prompt")
            mock_print.assert_called_once()
    
    @patch('requests.post')
    def test_call_llm_api_request_exception(self, mock_post):
        """Test API call with request exception"""
        mock_post.side_effect = requests.exceptions.RequestException("Network error")
        
        with patch('builtins.print') as mock_print:
            result = self.customizer.call_llm_api("test prompt")
            
            self.assertEqual(result, "test prompt")
            mock_print.assert_called_once()


class TestCVCustomizerSummarizeExperience(ParameterizedTestCase):
    """Test cases for CVCustomizer.summarize_experience method"""
    
    def setUp(self):
        """Set up test fixtures"""
        api_key = self.get_test_param('api_key', "test-key")
        api_endpoint = self.get_test_param('api_endpoint', "test-endpoint")
        config_file = self.get_test_param('config_file', "prompts_config.yaml")
        self.customizer = CVCustomizer(api_key, api_endpoint, config_file)
        
        self.sample_experience = [
            {
                "company": self.get_test_param('test_company_name', "Tech Corp"),
                "period": "2020-2023",
                "title": self.get_test_param('test_job_title', "Software Engineer"),
                "achievements": ["Built scalable systems", "Led team of 5"]
            }
        ]
        self.job_description = self.get_test_param('test_job_description', "Looking for a software engineer with leadership experience")
    
    @patch.object(CVCustomizer, 'call_llm_api')
    def test_summarize_experience_success(self, mock_call_llm):
        """Test successful experience summarization"""
        mock_call_llm.return_value = "Generated professional summary"
        
        result = self.customizer.summarize_experience(self.sample_experience, self.job_description)
        
        self.assertEqual(result, "Generated professional summary")
        mock_call_llm.assert_called_once()
        
        # Verify the prompt contains expected content
        call_args = mock_call_llm.call_args
        prompt = call_args[0][0]
        self.assertIn("Tech Corp", prompt)
        self.assertIn("Software Engineer", prompt)
        self.assertIn("Built scalable systems", prompt)
    
    @patch.object(CVCustomizer, 'call_llm_api')
    def test_summarize_experience_empty_achievements(self, mock_call_llm):
        """Test summarization with experience having no achievements"""
        experience_no_achievements = [
            {
                "company": "Tech Corp",
                "period": "2020-2023",
                "title": "Software Engineer"
            }
        ]
        mock_call_llm.return_value = "Generated summary"
        
        result = self.customizer.summarize_experience(experience_no_achievements, self.job_description)
        
        self.assertEqual(result, "Generated summary")
        mock_call_llm.assert_called_once()


class TestCVCustomizerCustomizeSummary(ParameterizedTestCase):
    """Test cases for CVCustomizer.customize_summary method"""
    
    def setUp(self):
        """Set up test fixtures"""
        api_key = self.get_test_param('api_key', "test-key")
        api_endpoint = self.get_test_param('api_endpoint', "test-endpoint")
        config_file = self.get_test_param('config_file', "prompts_config.yaml")
        self.customizer = CVCustomizer(api_key, api_endpoint, config_file)
        
        self.original_summary = "Experienced software engineer with 5 years of experience"
        self.job_description = self.get_test_param('test_job_description', "Looking for a senior developer with Python expertise")
    
    @patch.object(CVCustomizer, 'call_llm_api')
    def test_customize_summary_success(self, mock_call_llm):
        """Test successful summary customization"""
        mock_call_llm.return_value = "Customized summary for Python role"
        
        result = self.customizer.customize_summary(self.original_summary, self.job_description)
        
        self.assertEqual(result, "Customized summary for Python role")
        mock_call_llm.assert_called_once()
        
        # Verify the prompt contains expected content
        call_args = mock_call_llm.call_args
        prompt = call_args[0][0]
        self.assertIn(self.original_summary, prompt)
        self.assertIn(self.job_description, prompt)


class TestCVCustomizerCustomizeAchievements(ParameterizedTestCase):
    """Test cases for CVCustomizer.customize_achievements method"""
    
    def setUp(self):
        """Set up test fixtures"""
        api_key = self.get_test_param('api_key', "test-key")
        api_endpoint = self.get_test_param('api_endpoint', "test-endpoint")
        config_file = self.get_test_param('config_file', "prompts_config.yaml")
        self.customizer = CVCustomizer(api_key, api_endpoint, config_file)
        
        self.achievements = ["Built web applications", "Led development team"]
        self.job_description = self.get_test_param('test_job_description', "Looking for a full-stack developer with leadership skills")
        self.company_name = self.get_test_param('test_company_name', "Tech Corp")
        self.job_title = self.get_test_param('test_job_title', "Senior Developer")
    

    
    @patch.object(CVCustomizer, 'call_llm_api')
    def test_customize_achievements_success(self, mock_call_llm):
        """Test successful achievements customization"""
        mock_response = "- Developed scalable web applications\n- Managed cross-functional development team"
        mock_call_llm.return_value = mock_response
        
        result = self.customizer.customize_achievements(
            self.achievements, self.job_description, self.company_name, self.job_title
        )
        
        expected = ["Developed scalable web applications", "Managed cross-functional development team"]
        self.assertEqual(result, expected)
        mock_call_llm.assert_called_once()
    




class TestCVCustomizerCreateWordDocument(ParameterizedTestCase):
    """Test cases for CVCustomizer.create_word_document method"""
    
    def setUp(self):
        """Set up test fixtures"""
        api_key = self.get_test_param('api_key', "test-key")
        api_endpoint = self.get_test_param('api_endpoint', "test-endpoint")
        config_file = self.get_test_param('config_file', "prompts_config.yaml")
        self.customizer = CVCustomizer(api_key, api_endpoint, config_file)
        
        self.sample_cv_data = {
            "name": self.get_test_param('test_name', "John Doe"),
            "contact": {
                "phone": self.get_test_param('test_phone', "123-456-7890"),
                "email": self.get_test_param('test_email', "john@example.com"),
                "linkedin": self.get_test_param('test_linkedin', "linkedin.com/in/johndoe"),
                "nationality": self.get_test_param('test_nationality', "American")
            },
            "summary": "Experienced software engineer with 5 years of experience",
            "experience": [
                {
                    "company": self.get_test_param('test_company_name', "Tech Corp"),
                    "period": "2020-2023",
                    "title": self.get_test_param('test_job_title', "Software Engineer"),
                    "achievements": ["Built scalable systems", "Led team of 5"]
                }
            ],
            "qualifications": [
                {"year": "2022", "certification": "AWS Certified Developer"},
                {"year": "2021", "certification": "Python Certification"}
            ]
        }
    
    def test_create_word_document_success(self):
        """Test successful Word document creation"""
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
            temp_file.close()
            
            try:
                with patch('builtins.print') as mock_print:
                    self.customizer.create_word_document(self.sample_cv_data, temp_file.name)
                    
                    # Verify file was created
                    self.assertTrue(os.path.exists(temp_file.name))
                    
                    # Verify document content
                    doc = Document(temp_file.name)
                    self.assertEqual(doc.paragraphs[0].text, "John Doe")
                    self.assertIn("Summary", [p.text for p in doc.paragraphs])
                    self.assertIn("Professional Experience", [p.text for p in doc.paragraphs])
                    
                    # Save a copy in local directory for debugging (without printing)
                    self.save_temp_file_copy(temp_file.name, "output_success", verbose=False)
                    
                    mock_print.assert_called_once()
            finally:
                if os.path.exists(temp_file.name):
                    os.unlink(temp_file.name)
    
class TestCVCustomizerCustomizeCV(ParameterizedTestCase):
    """Test cases for CVCustomizer.customize_cv method"""
    
    def setUp(self):
        """Set up test fixtures"""
        api_key = self.get_test_param('api_key', "test-key")
        api_endpoint = self.get_test_param('api_endpoint', "test-endpoint")
        config_file = self.get_test_param('config_file', "prompts_config.yaml")
        self.customizer = CVCustomizer(api_key, api_endpoint, config_file)
        
        self.sample_cv_data = {
            "name": self.get_test_param('test_name', "John Doe"),
            "contact": {"email": self.get_test_param('test_email', "john@example.com")},
            "experience": [
                {
                    "company": self.get_test_param('test_company_name', "Tech Corp"),
                    "period": "2020-2023",
                    "title": self.get_test_param('test_job_title', "Software Engineer"),
                    "achievements": ["Built scalable systems"]
                }
            ]
        }
        self.job_description = self.get_test_param('test_job_description', "Looking for a software engineer")
    
    @patch.object(CVCustomizer, 'load_cv_template')
    @patch.object(CVCustomizer, 'customize_achievements')
    @patch.object(CVCustomizer, 'summarize_experience')
    @patch.object(CVCustomizer, 'create_word_document')
    def test_customize_cv_success(self, mock_create_doc, mock_summarize, mock_customize_achievements, mock_load_template):
        """Test successful CV customization"""
        mock_load_template.return_value = self.sample_cv_data
        mock_customize_achievements.return_value = ["Customized achievement"]
        mock_summarize.return_value = "Customized summary"
        
        with patch('builtins.print') as mock_print:
            result = self.customizer.customize_cv("template.yaml", self.job_description, "output.docx")
            
            self.assertEqual(result, "output.docx")
            
            # Verify all methods were called
            mock_load_template.assert_called_once_with("template.yaml")
            mock_customize_achievements.assert_called_once()
            mock_summarize.assert_called_once()
            mock_create_doc.assert_called_once()
            
            # Verify print statements
            self.assertEqual(mock_print.call_count, 4)
    
    @patch.object(CVCustomizer, 'load_cv_template')
    @patch.object(CVCustomizer, 'customize_achievements')
    @patch.object(CVCustomizer, 'summarize_experience')
    @patch.object(CVCustomizer, 'create_word_document')
    def test_customize_cv_without_output_file(self, mock_create_doc, mock_summarize, mock_customize_achievements, mock_load_template):
        """Test CV customization without specifying output file"""
        mock_load_template.return_value = self.sample_cv_data
        mock_customize_achievements.return_value = ["Customized achievement"]
        mock_summarize.return_value = "Customized summary"
        
        with patch('datetime.datetime') as mock_datetime:
            mock_datetime.now.return_value.strftime.return_value = "20240101_120000"
            
            result = self.customizer.customize_cv("template.yaml", self.job_description)
            
            self.assertTrue(result.startswith("customized_cv_"))
            self.assertTrue(result.endswith(".docx"))
            mock_create_doc.assert_called_once()
    
    @patch.object(CVCustomizer, 'load_cv_template')
    def test_customize_cv_template_loading_error(self, mock_load_template):
        """Test CV customization with template loading error"""
        mock_load_template.side_effect = FileNotFoundError("Template not found")
        
        with self.assertRaises(FileNotFoundError):
            self.customizer.customize_cv("nonexistent.yaml", self.job_description)


def run_individual_test(test_class, test_method=None):
    """Run an individual test class or specific test method"""
    suite = unittest.TestSuite()
    
    if test_method:
        suite.addTest(test_class(test_method))
    else:
        suite.addTests(unittest.defaultTestLoader.loadTestsFromTestCase(test_class))
    
    runner = unittest.TextTestRunner(verbosity=2)
    result = runner.run(suite)
    return result.wasSuccessful()


def main():
    """Main function to run all tests or specific test categories"""
    import sys
    
    # Parse command line arguments
    parser = argparse.ArgumentParser(description='CVCustomizer Unit Tests with Parameters')
    parser.add_argument('test_name', nargs='?', help='Specific test category to run')
    parser.add_argument('--api-key', help='API key for testing')
    parser.add_argument('--api-endpoint', help='API endpoint for testing')
    parser.add_argument('--test-yaml-file', help='YAML file for testing')
    parser.add_argument('--test-output-file', help='Output file for testing')
    parser.add_argument('--test-job-description', help='Job description for testing')
    parser.add_argument('--test-company-name', help='Company name for testing')
    parser.add_argument('--test-job-title', help='Job title for testing')
    parser.add_argument('--test-name', help='Name for testing')
    parser.add_argument('--test-email', help='Email for testing')
    parser.add_argument('--test-phone', help='Phone for testing')
    parser.add_argument('--test-linkedin', help='LinkedIn for testing')
    parser.add_argument('--test-nationality', help='Nationality for testing')
    
    args = parser.parse_args()
    
    if args.test_name:
        # Map test names to test classes
        test_classes = {
            'init': TestCVCustomizerInit,
            'load_template': TestCVCustomizerLoadCVTemplate,
            'call_llm_api': TestCVCustomizerCallLLMAPI,
            'summarize_experience': TestCVCustomizerSummarizeExperience,
            'customize_summary': TestCVCustomizerCustomizeSummary,
            'customize_achievements': TestCVCustomizerCustomizeAchievements,
            'create_word_document': TestCVCustomizerCreateWordDocument,
            'customize_cv': TestCVCustomizerCustomizeCV
        }
        
        if args.test_name in test_classes:
            print(f"Running {args.test_name} tests with parameters...")
            success = run_individual_test(test_classes[args.test_name])
            sys.exit(0 if success else 1)
        else:
            print(f"Unknown test: {args.test_name}")
            print("Available tests:", list(test_classes.keys()))
            sys.exit(1)
    else:
        # Run all tests
        print("Running all tests with parameters...")
        unittest.main(verbosity=2)


if __name__ == "__main__":
    main()

