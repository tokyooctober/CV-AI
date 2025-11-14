#!/usr/bin/env python3
"""
CV Customization System
Creates customized CVs based on job descriptions using LLM API for content adaptation.
"""

import yaml
import requests
import json
import argparse
import os
from docx.shared import Pt
from datetime import datetime
from docx import Document
from docx.shared import Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import re

#OPENAI_API_KEY = "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"
OPENAI_API_KEY = "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"
OPENAI_API_ENDPOINT = "https://api.openai.com/v1/chat/completions"


class CVCustomizer:

    def __init__(self, api_key=OPENAI_API_KEY, api_endpoint=OPENAI_API_ENDPOINT, config_file="prompts_config.yaml"):
        """
        Initialize the CV Customizer with API credentials and configuration.

        Args:
            api_key (str): API key for the LLM service
            api_endpoint (str): API endpoint URL
            config_file (str): Path to configuration file
        """
        self.api_key = api_key or os.getenv("LLM_API_KEY")
        self.api_endpoint = api_endpoint or os.getenv(
            "LLM_API_ENDPOINT", "https://api.openai.com/v1/chatP/completions"
        )
        
        # Load configuration
        self.config = self._load_config(config_file)

        if not self.api_key:
            print(self.config['messages']['warnings']['no_api_key'])
        
        # Will store the response id from an initial JD context call
        self.llm_context_id = None
    
    def _load_config(self, config_file):
        """Load configuration from YAML file"""
        try:
            with open(config_file, "r", encoding="utf-8") as file:
                return yaml.safe_load(file)
        except FileNotFoundError:
            print(f"Warning: Configuration file '{config_file}' not found. Using default values.")
            return self._get_default_config()
        except yaml.YAMLError as e:
            print(f"Warning: Error parsing configuration file: {e}. Using default values.")
            return self._get_default_config()
    
    def _get_default_config(self):
        """Get default configuration as fallback"""
        return {
            'messages': {
                'warnings': {
                    'no_api_key': "Warning: No API key provided. Set LLM_API_KEY environment variable or pass api_key parameter.",
                    'no_api_key_available': "Warning: No API key available. Returning original content.",
                    'unexpected_api_response': "Warning: Unexpected API response format. Returning original content.",
                    'api_call_failed': "Warning: API call failed: {error}. Returning original content."
                }
            }
        }

    def load_cv_template(self, yaml_file):
        """
        Load CV template from YAML file.

        Args:
            yaml_file (str): Path to YAML template file

        Returns:
            dict: CV template data
        """
        try:
            with open(yaml_file, "r", encoding="utf-8") as file:
                return yaml.safe_load(file)
        except FileNotFoundError:
            error_msg = self.config.get('messages', {}).get('errors', {}).get('file_not_found', "CV template file '{file}' not found.")
            raise FileNotFoundError(error_msg.format(file=yaml_file))
        except yaml.YAMLError as e:
            error_msg = self.config.get('messages', {}).get('errors', {}).get('yaml_parse_error', "Error parsing YAML file: {error}")
            raise ValueError(error_msg.format(error=e))

    def call_llm_api(self, prompt, system_message=None):
        """
        Make API call to LLM service for content customization.

        Args:
            prompt (str): The prompt to send to the LLM
            system_message (str): Optional system message

        Returns:
            str: LLM response
        """
        if not self.api_key:
            print(self.config['messages']['warnings']['no_api_key_available'])
            return prompt

        # Get API configuration
        api_config = self.config.get('api', {})
        headers_config = self.config.get('http_headers', {})
        
        headers = {
            headers_config.get('authorization', 'Authorization'): f"{headers_config.get('bearer_prefix', 'Bearer {api_key}').format(api_key=self.api_key)}",
            headers_config.get('content_type', 'Content-Type'): headers_config.get('application_json', 'application/json'),
        }

        messages = []
        if system_message:
            messages.append({"role": "system", "content": system_message})

        messages.append({"role": "user", "content": prompt})

        #print(f"messages: {messages}")

        data = {
            "model": api_config.get('model', 'gpt-3.5-turbo'),
            "messages": messages,
            "temperature": api_config.get('temperature', 0.7),
            "max_tokens": api_config.get('max_tokens', 2000),
        }

        try:
            response = requests.post(
                self.api_endpoint, headers=headers, json=data, timeout=api_config.get('timeout', 30)
            )
            response.raise_for_status()

            result = response.json()
            api_response_config = self.config.get('api_response', {})
            choices_key = api_response_config.get('choices_key', 'choices')
            message_key = api_response_config.get('message_key', 'message')
            content_key = api_response_config.get('content_key', 'content')
            min_choices = api_response_config.get('min_choices_length', 1)
            
            if choices_key in result and len(result[choices_key]) >= min_choices:
                return result[choices_key][0][message_key][content_key].strip()
            else:
                print(self.config['messages']['warnings']['unexpected_api_response'])
                return prompt

        except requests.exceptions.RequestException as e:
            error_msg = self.config['messages']['warnings']['api_call_failed']
            print(error_msg.format(error=e))
            return prompt

    # def create_jd_context(self, job_description):
    #     """
    #     Create an initial LLM context using the job description and capture the response id
    #     for use in subsequent calls.

    #     Args:
    #         job_description (str): Job description text

    #     Returns:
    #         str | None: Response id from the LLM call, or None if unavailable
    #     """
    #     if not self.api_key:
    #         print(self.config['messages']['warnings']['no_api_key_available'])
    #         return None

    #     api_config = self.config.get('api', {})
    #     headers_config = self.config.get('http_headers', {})

    #     headers = {
    #         headers_config.get('authorization', 'Authorization'): f"{headers_config.get('bearer_prefix', 'Bearer {api_key}').format(api_key=self.api_key)}",
    #         headers_config.get('content_type', 'Content-Type'): headers_config.get('application_json', 'application/json'),
    #     }

    #     # Use a lightweight system instruction to establish context
    #     messages = [
    #         {"role": "system", "content": "Remember the provided Job Description as context for future chats."},
    #         {"role": "user", "content": f"Job Description (store as context):\n{job_description}"},
    #     ]

    #     data = {
    #         "model": api_config.get('model', 'gpt-3.5-turbo'),
    #         "messages": messages,
    #         "temperature": api_config.get('temperature', 0.2),
    #         "max_tokens": api_config.get('max_tokens', 64),
    #     }

    #     try:
    #         response = requests.post(
    #             self.api_endpoint, headers=headers, json=data, timeout=api_config.get('timeout', 30)
    #         )
    #         response.raise_for_status()
    #         result = response.json()
    #         print(f"create_jd_context response: {result}")
    #         # Root-level id is expected for OpenAI chat completions
    #         context_id = result.get('id')
    #         print(f"LLM context initialized with response id: {context_id}")    
    #         self.llm_context_id = context_id
    #         return context_id
    #     except requests.exceptions.RequestException as e:
    #         error_msg = self.config['messages']['warnings']['api_call_failed']
    #         print(error_msg.format(error=e))
    #         return None

    def summarize_experience(self, job_description, experience_data):
        """
        Generate a professional summary based on experience data and job description.

        Args:
            experience_data (list): List of experience dictionaries
            job_description (str): Job description to match

        Returns:
            str: Generated professional summary
        """
        system_message = self.config['system_messages']['summarize_experience']
        #system_message = system_template.format(job_description=job_description)

        # Format experience data for the prompt
        experience_templates = self.config.get('experience_templates', {})
        experience_text = ""
        for exp in experience_data:
            experience_text += f"\n{experience_templates.get('company', 'Company: {company}').format(company=exp['company'])}\n"
            experience_text += f"{experience_templates.get('period', 'Period: {period}').format(period=exp['period'])}\n"
            experience_text += f"{experience_templates.get('title', 'Title: {title}').format(title=exp['title'])}\n"
            if exp.get('achievements'):
                experience_text += f"{experience_templates.get('achievements_header', 'Key Achievements:')}\n"
                for achievement in exp['achievements']:
                    experience_text += f"{experience_templates.get('achievement_item', '- {achievement}').format(achievement=achievement)}\n"
            experience_text += "\n"

        prompt_template = self.config['prompts']['summarize_experience']
        prompt = prompt_template.format(
            job_description=job_description,
            experience_text=experience_text)

        return self.call_llm_api(prompt, system_message)

    # def customize_summary(self, original_summary, job_description):
    #     """
    #     Customize CV summary based on job description.

    #     Args:
    #         original_summary (str): Original CV summary
    #         job_description (str): Job description to match

    #     Returns:
    #         str: Customized summary
    #     """
    #     system_message = self.config['system_messages']['customize_summary']

    #     prompt_template = self.config['prompts']['customize_summary']
    #     prompt = prompt_template.format(
    #         job_description=job_description,
    #         original_summary=original_summary
    #     )

    #     return self.call_llm_api(prompt, system_message)

    def customize_achievements(
        self, job_description, achievements, company_name, job_title
    ):
        """
        Customize achievements for a specific role based on job description.

        Args:
            achievements (list): List of achievement strings
            job_description (str): Job description to match
            company_name (str): Company name for context
            job_title (str): Job title for context

        Returns:
            list: Customized achievements
        """
        if not achievements:
            return []

        system_message = self.config['system_messages']['customize_achievements']

        achievement_patterns = self.config.get('achievement_patterns', {})
        dash_prefix = achievement_patterns.get('dash_prefix', '- ')

        achievements_text = "\n".join(
            [f"{dash_prefix}{achievement}" for achievement in achievements]
        )

        prompt_template = self.config['prompts']['customize_achievements']
        prompt = prompt_template.format(
            job_description=job_description,
            company_name=company_name,
            job_title=job_title,
            job_description=job_description,
            achievements_text=achievements_text
        )

        response = self.call_llm_api(prompt, system_message)

        # Parse the response back into a list
        customized_achievements = []
        dash_prefix = achievement_patterns.get('dash_prefix', '- ')
        bullet_prefix = achievement_patterns.get('bullet_prefix', '• ')
        comment_prefix = achievement_patterns.get('comment_prefix', '#')
        
        for line in response.split("\n"):
            line = line.strip()
            if line.startswith(dash_prefix):
                customized_achievements.append(line[len(dash_prefix):])
            elif line.startswith(bullet_prefix):
                customized_achievements.append(line[len(bullet_prefix):])
            elif line and not line.startswith(comment_prefix):
                customized_achievements.append(line)

        return customized_achievements if customized_achievements else achievements

    def add_title(self, doc, text):
        """
        Add title to the document.

        Args:
            doc (Document): Document object
            text (str): Title text

        Returns:
            str: Title text
        """
        # Create a single cell table
        # Create a new Document and add a single-cell table
        table = doc.add_table(rows=1, cols=1)
        cell = table.cell(0, 0)

        # Remove left and right borders
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        tcBorders = OxmlElement('w:tcBorders')
        for border_name in ('top', 'bottom', 'left', 'right'):
            border = OxmlElement(f'w:{border_name}')
            if border_name in ('left', 'right'):
                border.set(qn('w:val'), 'nil')
            else:
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), 'auto')
            tcBorders.append(border)
        tcPr.append(tcBorders)

        cell.text = ""
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Remove all top and bottom spacings in the paragraph
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(text)
        run.bold = True
        run.font.size = Pt(12)
        run.font.name = 'Calibri'

        return table

        
    def create_word_document(self, cv_data, output_file):
        """
        Create Word document with the customized CV.

        Args:
            cv_data (dict): CV data with customized content
            output_file (str): Output file path
        """
        doc = Document()

        # Set default paragraph spacing to 0 before and after for the whole document
        # Set global default paragraph spacing to 0 before and after for all styles
        for style in doc.styles:
            if style.type == 1:  # 1 = WD_STYLE_TYPE.PARAGRAPH
                if hasattr(style, 'paragraph_format'):
                    style.paragraph_format.space_before = Pt(0)
                    style.paragraph_format.space_after = Pt(0)
                if hasattr(style, 'font'):
                    style.font.name = 'Calibri'
                    style.font.size = Pt(10)

        # Set up document margins
        doc_sections = doc.sections
        for section in doc_sections:
            section.top_margin = Cm(1.27)
            section.bottom_margin = Cm(1.27)
            section.left_margin = Cm(1.27)
            section.right_margin = Cm(1.27)
            # Set A4 page size for all sections
            section.page_width = Cm(21.0)
            section.page_height = Cm(29.7)

        # Add name as title
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run(cv_data["name"])
        title_run.bold = True
        title_run.font.size = Pt(16)
        #title_run.font.name = 'Calibri'

        # Add contact information
        contact = cv_data["contact"]
        contact_para = doc.add_paragraph()
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        contact_info = []
        if "phone" in contact:
            contact_info.append(contact["phone"])
        if "email" in contact:
            contact_info.append(contact["email"])
        if "linkedin" in contact:
            contact_info.append(contact["linkedin"])
        if "nationality" in contact:
            contact_info.append(contact["nationality"])

        contact_separator = " | "
        contact_para.add_run(contact_separator.join(contact_info))

        # doc_strings = self.config.get('document_strings', {})
        # sections = doc_strings.get('sections', {})

        self.add_title(doc, "SUMMARY")

        # Add summary section
        summary_para = doc.add_paragraph()
        summary_para.add_run(cv_data["summary"])

        # Add professional experience section
        self.add_title(doc, "PROFESSIONAL EXPERIENCE")

        for experience in cv_data["experience"]:
            # Company name and period (left and right justified)
            company_para = doc.add_paragraph()
        
            company_para.paragraph_format.tab_stops.clear_all()
            company_para.paragraph_format.tab_stops.add_tab_stop(Inches(7), WD_ALIGN_PARAGRAPH.RIGHT)

            run = company_para.add_run(experience["company"]+"\t" + experience["period"])
            run.font.size = Pt(12)
            run.bold = True
 
            # Job title
            title_para = doc.add_paragraph()
            run = title_para.add_run(experience["title"])
            run.font.size = Pt(10)
            run.bold = True
            title_para.paragraph_format.left_indent = 0

            # Achievements
            if experience["achievements"]:
                for achievement in experience["achievements"]:
                    achievement_para = doc.add_paragraph()
                    tab_stop_width = Inches(0.24)
                    achievement_para.paragraph_format.left_indent = tab_stop_width
                    achievement_para.paragraph_format.first_line_indent = -tab_stop_width
                    
                    tab_stops = achievement_para.paragraph_format.tab_stops
                    tab_stops.add_tab_stop(tab_stop_width, WD_TAB_ALIGNMENT.LEFT)

                    achievement_para.add_run("•\t" + achievement)

            # Add space between experiences
            doc.add_paragraph()

        # Add certifications section
        if "qualifications" in cv_data and cv_data["qualifications"]:

            self.add_title(doc, "QUALIFICATIONS & PROFESSIONAL DEVELOPMENT")
            
            # Add certification data
            for qualification in cv_data["qualifications"]:
                year = str(qualification.get("year", ""))
                certification = qualification.get("certification", "")

                para = doc.add_paragraph()
                pf = para.paragraph_format

                # define the tab column
                tab_width = Inches(1)
                pf.tab_stops.add_tab_stop(tab_width, WD_TAB_ALIGNMENT.LEFT)

                pf.left_indent = tab_width
                pf.first_line_indent = -tab_width

                # text with a tab character
                if year:
                    text = f"{year}\t{certification}"
                else:
                    text = f"\t{certification}"  # <-- tab even when year is empty

                para.add_run(text)
     

        # Save the document
        doc.save(output_file)
        success_msg = self.config.get('messages', {}).get('success', {}).get('cv_saved', 'Customized CV saved to: {output_file}')
        print(success_msg.format(output_file=output_file))

    def customize_cv(self, yaml_file, job_description, output_file=None):
        """
        Main method to customize CV based on job description.

        Args:
            yaml_file (str): Path to YAML template file
            job_description (str): Job description text
            output_file (str): Output file path (optional)

        Returns:
            str: Path to generated Word document
        """
        # Load CV template
        step_messages = self.config.get('step_messages', {})
        print(step_messages.get('loading_template', 'Step 1: Loading CV template...'))
        cv_data = self.load_cv_template(yaml_file)

        #HERE
        # Initialize LLM context with the Job Description and store the response id
        # print(step_messages.get('initializing_llm_context', 'Step 1b: Initializing LLM context from Job Description...'))
        # self.create_jd_context(job_description)

        # Customize achievements for each experience
        print(step_messages.get('customizing_achievements', 'Step 2: Customizing achievements...'))
        for experience in cv_data["experience"]:
            if experience["achievements"]:
                # Print company and title before customizing achievements
                print(f"Customizing achievements for: {experience['company']} - {experience['title']}")
                print(f"Length before customizing: {len(experience['achievements'])}")
                experience["achievements"] = self.customize_achievements(
                    job_description,
                    experience["achievements"],
                    job_description,
                    experience["company"],
                    experience["title"],
                )
                print(f"Length after customizing: {len(experience['achievements'])}")

        # Generate summary from experience
        print(step_messages.get('generating_summary', 'Step 3: Generating summary from experience...'))
        cv_data["summary"] = self.summarize_experience(job_description, cv_data["experience"])

        # Generate output filename if not provided
        if not output_file:
            output_settings = self.config.get('output_settings', {})
            timestamp = datetime.now().strftime(output_settings.get('timestamp_format', '%Y%m%d_%H%M%S'))
            yaml_prefix = os.path.splitext(os.path.basename(yaml_file))[0]
            # Try to get the job description filename without extension, or fallback to string representation
            if os.path.isfile(job_description):
                jd_prefix = os.path.splitext(os.path.basename(job_description))[0]
            else:
                jd_prefix = os.path.splitext(os.path.basename(str(job_description)))[0]
            prefix = f"{yaml_prefix}_{jd_prefix}_"
            extension = output_settings.get('file_extension', '.docx')
            output_file = f"{prefix}{timestamp}{extension}"  

        # Create Word document
        print(step_messages.get('generating_document', 'Step 4: Generating Word document...'))
        self.create_word_document(cv_data, output_file)

        return output_file


def main():
    """Main function to run the CV customizer."""
    parser = argparse.ArgumentParser(
        description="Customize CV based on job description"
    )
    parser.add_argument("yaml_file", help="Path to YAML CV template file")
    parser.add_argument("job_description", help="Job description text or path to file")
    # parser.add_argument("--api-key", help="LLM API key")
    # parser.add_argument("--api-endpoint", help="LLM API endpoint")

    args = parser.parse_args()
    api_key = OPENAI_API_KEY
    api_endpoint = OPENAI_API_ENDPOINT 
    # Read job description from file if it's a file path
    job_description = args.job_description
    if os.path.isfile(job_description):
        with open(job_description, "r", encoding="utf-8") as f:
            job_description = f.read()
    else:
        # Load configuration for error message
        temp_customizer = CVCustomizer(api_key, api_endpoint)
        error_msg = temp_customizer.config.get('messages', {}).get('errors', {}).get('job_description_file_error', 'Pls provide filename to the job desscription. Unable to open the file:{file}.')
        print(f"\n{error_msg.format(file=job_description)}")
        return  

    # Initialize customizer
    customizer = CVCustomizer(api_key, api_endpoint)

    # Customize CV
    output_file = customizer.customize_cv(args.yaml_file, job_description, output_file=None)

    success_messages = customizer.config.get('messages', {}).get('success', {})
    print(f"\n{success_messages.get('cv_customization_completed', 'CV customization completed successfully!')}")
    print(f"{success_messages.get('output_file', 'Output file: {output_file}').format(output_file=output_file)}")


if __name__ == "__main__":
    main()
